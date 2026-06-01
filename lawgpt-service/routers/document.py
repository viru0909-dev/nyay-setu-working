"""
routers/document.py — Legal document generation endpoints for Nyay Setu LawGPT.

Provides:
    POST /generate      — Generate legal document text (affidavit, RTI, complaint, notice)
    POST /generate/pdf  — Generate legal document as downloadable PDF
"""

import io
import logging
import os
import tempfile
from datetime import datetime
from typing import Dict, List, Literal, Optional

from dotenv import load_dotenv
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from lawgpt.retriever import retrieve
from lawgpt.prompt_builder import build_prompt, validate_required_fields, detect_prompt_injection

load_dotenv()
logger = logging.getLogger("lawgpt")

router = APIRouter()


# ── Request / Response models ──────────────────────────────────────────────────

class DocumentFields(BaseModel):
    petitioner_name: str
    petitioner_address: str
    respondent_name: Optional[str] = ""
    respondent_address: Optional[str] = ""
    case_description: str
    incident_date: str
    relief_sought: Optional[str] = ""
    court_name: Optional[str] = ""
    department_name: Optional[str] = ""
    pio_name: Optional[str] = ""


class GenerateRequest(BaseModel):
    doc_type: Literal["affidavit", "rti", "complaint", "notice"]
    fields: DocumentFields
    language: str = Field(default="en")


class GenerateResponse(BaseModel):
    doc_type: str
    title: str
    content: str
    legal_context: str
    sources: List[str]
    generated_at: str


# ── Prompt templates ───────────────────────────────────────────────────────────

AFFIDAVIT_PROMPT: str = """You are an expert Indian legal drafter. Generate a formal AFFIDAVIT.
Use this legal context for accuracy: {legal_context}

Petitioner: {petitioner_name}, {petitioner_address}
Respondent: {respondent_name}, {respondent_address}
Incident: {case_description}
Date: {incident_date}
Relief sought: {relief_sought}

Format EXACTLY as:
AFFIDAVIT

I, {petitioner_name}, son/daughter of _______, aged ___ years,
residing at {petitioner_address}, do hereby solemnly affirm and state as follows:

1. That I am the deponent herein and am fully conversant with the facts stated herein.
2. [Facts of the case in numbered paragraphs]
3. [Legal grounds]

PRAYER:
[Relief sought]

VERIFICATION:
Verified at _______ on this ___ day of _______, 20__.
That the contents of the above affidavit are true to the best of my knowledge.

Deponent
"""

RTI_PROMPT: str = """You are an expert Indian legal drafter. Generate a formal RTI APPLICATION
under the Right to Information Act, 2005.
Use this legal context: {legal_context}

Applicant: {petitioner_name}, {petitioner_address}
Department/Public Authority: {department_name}
PIO Name: {pio_name}
Information sought: {case_description}

Format EXACTLY as:
APPLICATION UNDER RIGHT TO INFORMATION ACT, 2005
Section 6(1)

To,
The Public Information Officer,
{department_name}

Subject: Request for Information under RTI Act, 2005

Sir/Madam,

I, {petitioner_name}, a citizen of India, hereby request the following
information under Section 6(1) of the Right to Information Act, 2005:

INFORMATION REQUESTED:
1. [Specific information point 1]
2. [Specific information point 2]
[Continue as needed]

I am willing to pay the prescribed fee. Please provide the information within
30 days as mandated under Section 7(1) of the RTI Act, 2005.

If the information is denied, please provide reasons under Section 8 of the Act.

Yours faithfully,
{petitioner_name}
{petitioner_address}
Date: {incident_date}
"""

COMPLAINT_PROMPT: str = """You are an expert Indian legal drafter. Generate a formal LEGAL COMPLAINT.
Use this legal context for accurate section references: {legal_context}

Complainant: {petitioner_name}, {petitioner_address}
Accused/Respondent: {respondent_name}, {respondent_address}
Incident: {case_description}
Date of incident: {incident_date}
Relief sought: {relief_sought}

Format EXACTLY as:
COMPLAINT UNDER SECTION ___ OF THE CODE OF CRIMINAL PROCEDURE, 1973

IN THE COURT OF THE JUDICIAL MAGISTRATE
[Court name]

COMPLAINT NO. ___/20__

Complainant: {petitioner_name}
s/o d/o _______, aged ___ years
Address: {petitioner_address}

Versus

Accused: {respondent_name}
Address: {respondent_address}

COMPLAINT

Most Respectfully Showeth:

1. That the Complainant is a citizen of India residing at the above address.
2. [Facts in numbered paragraphs with dates]
3. [Legal grounds citing relevant IPC/BNS sections from context]
4. [Specific offences committed]

PRAYER:
It is therefore most respectfully prayed that this Hon'ble Court may be
pleased to:
[Specific relief sought]

Place: _______
Date: {incident_date}

Complainant
{petitioner_name}
"""

NOTICE_PROMPT: str = """You are an expert Indian legal drafter. Generate a formal LEGAL NOTICE.
Use this legal context: {legal_context}

Sender: {petitioner_name}, {petitioner_address}
Recipient: {respondent_name}, {respondent_address}
Issue: {case_description}
Date: {incident_date}
Relief/Demand: {relief_sought}

Format EXACTLY as:
LEGAL NOTICE

Date: {incident_date}

To,
{respondent_name}
{respondent_address}

TAKE NOTICE that my client {petitioner_name} of {petitioner_address}
has instructed me to serve upon you this legal notice as under:

1. That my client is [relationship/context].
2. [Facts in numbered paragraphs]
3. [Legal provisions violated — cite from context]
4. [Demand/Relief]

You are hereby called upon to [specific action] within 15 days of
receipt of this notice, failing which my client shall be constrained
to initiate appropriate legal proceedings against you without further
notice, at your risk, cost and consequences.

[Advocate Name]
Advocate
[Address]
"""

PROMPT_MAP: Dict[str, str] = {
    "affidavit": AFFIDAVIT_PROMPT,
    "rti": RTI_PROMPT,
    "complaint": COMPLAINT_PROMPT,
    "notice": NOTICE_PROMPT,
}

TITLE_MAP: Dict[str, str] = {
    "affidavit": "AFFIDAVIT",
    "rti": "APPLICATION UNDER RIGHT TO INFORMATION ACT, 2005",
    "complaint": "LEGAL COMPLAINT",
    "notice": "LEGAL NOTICE",
}


# ── External templates loader (optional) ────────────────────────────────────────
import json
from pathlib import Path

# Module-level templates dict (may be empty if no external JSON provided)
_templates: Dict[str, dict] = {}

_templates_path = Path(__file__).parent.parent / "templates" / "document_templates.json"
if _templates_path.exists():
    try:
        with _templates_path.open("r", encoding="utf-8") as fh:
            _templates = json.load(fh)
        # Override PROMPT_MAP and TITLE_MAP if keys exist in JSON
        for k, v in _templates.items():
            if "prompt_template" in v:
                PROMPT_MAP[k] = v["prompt_template"]
            if "title" in v:
                TITLE_MAP[k] = v["title"]
        logger.info("Loaded document templates from %s", _templates_path)
    except Exception:
        logger.exception("Failed to load external document templates — using defaults")
else:
    logger.debug("No external document templates found at %s", _templates_path)


# ── LLM resolution (shared with context.py) ───────────────────────────────────

_doc_llm = None
_doc_llm_label: str = "none"


def _get_doc_llm():
    """Lazy-initialise and return the LLM instance for document generation.

    Falls back to a lightweight DummyLLM when langchain or other LLM
    integrations are not available (useful for local QA and tests).
    """
    global _doc_llm, _doc_llm_label

    if _doc_llm is not None:
        return _doc_llm, _doc_llm_label

    groq_key: Optional[str] = os.getenv("GROQ_API_KEY")
    gemini_key: Optional[str] = os.getenv("GEMINI_API_KEY")

    try:
        if groq_key:
            from langchain_groq import ChatGroq
            _doc_llm = ChatGroq(
                model="llama-3.3-70b-versatile",
                temperature=0.3,
                groq_api_key=groq_key,
                max_tokens=2000,
            )
            _doc_llm_label = "groq"
            logger.info("📝 Document LLM: Groq (llama-3.3-70b-versatile)")
        elif gemini_key:
            from langchain_google_genai import ChatGoogleGenerativeAI
            _doc_llm = ChatGoogleGenerativeAI(
                model="gemini-1.5-pro",
                temperature=0.3,
                google_api_key=gemini_key,
                max_output_tokens=2000,
            )
            _doc_llm_label = "gemini"
            logger.info("📝 Document LLM: Google Gemini (gemini-1.5-pro)")
        else:
            from langchain_community.llms import Ollama
            _doc_llm = Ollama(
                model="llama3",
                base_url="http://localhost:11434",
                temperature=0.3,
            )
            _doc_llm_label = "ollama"
            logger.info("📝 Document LLM: Ollama (llama3, local)")

        return _doc_llm, _doc_llm_label

    except Exception as e:
        fake_llm_flag = os.getenv("LAWGPT_FAKE_LLM") == "1"
        if fake_llm_flag:
            class DummyLLM:
                def invoke(self, prompt):
                    return {
                        "content": f"DUMMY GENERATED DOCUMENT\n\n{prompt[:400]}"
                    }

            _doc_llm = DummyLLM()
            _doc_llm_label = "dummy"
            logger.warning("Using DummyLLM fallback because LLM integrations are not available: %s", e)
            return _doc_llm, _doc_llm_label

        logger.error("No LLM integration available and LAWGPT_FAKE_LLM is not enabled: %s", e)
        raise


# ── Core generation logic ─────────────────────────────────────────────────────

def _generate_document(request: GenerateRequest) -> GenerateResponse:
    """Shared generation logic used by both /generate and /generate/pdf."""
    fields: DocumentFields = request.fields
    doc_type: str = request.doc_type

    # 1. Retrieve legal context from FAISS
    search_query: str = f"{doc_type} {fields.case_description}"
    try:
        results = retrieve(query=search_query, k=3)
    except FileNotFoundError:
        raise HTTPException(
            status_code=503,
            detail="Legal database not initialized. Run 'python lawgpt/ingest.py' first.",
        )
    except ImportError:
        # langchain/FAISS not available in this environment (tests/dev). Continue with empty context.
        logger.warning("langchain_community not available; proceeding without legal context")
        results = []

    # Build context and sources
    context_parts: List[str] = []
    sources: List[str] = []
    for doc in results:
        context_parts.append(f"- {doc.page_content}")
        source_name: str = doc.metadata.get("source", "unknown")
        page_num: int = doc.metadata.get("page", 0)
        source_label: str = f"{source_name} — page {page_num}"
        if source_label not in sources:
            sources.append(source_label)

    legal_context: str = "\n\n".join(context_parts) if context_parts else "No specific legal context available."

    # 2. Build prompt
    prompt_template: str = PROMPT_MAP[doc_type]
    # Prepare a simple dict of fields for prompt construction
    field_map = {
        "petitioner_name": fields.petitioner_name,
        "petitioner_address": fields.petitioner_address,
        "respondent_name": fields.respondent_name or "N/A",
        "respondent_address": fields.respondent_address or "N/A",
        "case_description": fields.case_description,
        "incident_date": fields.incident_date,
        "relief_sought": fields.relief_sought or "N/A",
        "court_name": fields.court_name or "___",
        "department_name": fields.department_name or "___",
        "pio_name": fields.pio_name or "The Public Information Officer",
    }

    # If external templates provided required_fields, enforce them
    required = []
    if _templates and doc_type in _templates:
        required = _templates[doc_type].get("required_fields", [])

    # Validate required fields early and return a clear error to the client
    missing = validate_required_fields({k: v for k, v in field_map.items()}, required)
    if missing:
        raise HTTPException(status_code=422, detail={"missing_fields": missing})

    # Check for obvious prompt-injection patterns in user inputs
    suspicious = detect_prompt_injection({k: v for k, v in field_map.items()})
    if suspicious:
        raise HTTPException(status_code=400, detail={"prompt_injection_detected": suspicious})

    prompt: str = build_prompt(prompt_template, field_map, legal_context=legal_context)

    # 3. Call LLM
    try:
        llm, label = _get_doc_llm()
        answer = llm.invoke(prompt)
        if hasattr(answer, "content"):
            answer = answer.content
        content: str = str(answer)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Document LLM error: %s", e, exc_info=True)
        raise HTTPException(status_code=502, detail=f"LLM error: {e}")

    logger.info("📝 Generated %s for %s", doc_type, fields.petitioner_name)

    return GenerateResponse(
        doc_type=doc_type,
        title=TITLE_MAP[doc_type],
        content=content,
        legal_context=legal_context,
        sources=sources,
        generated_at=datetime.utcnow().isoformat() + "Z",
    )


# ── PDF generation helper ─────────────────────────────────────────────────────

def _create_pdf(response: GenerateResponse, petitioner_name: str) -> io.BytesIO:
    """Convert generated document text to a styled A4 PDF using ReportLab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="reportlab is not installed. Run: pip install reportlab==4.2.2",
        )

    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    # Styles
    styles = getSampleStyleSheet()

    header_style = ParagraphStyle(
        "DocHeader",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        textColor="#666666",
        alignment=TA_CENTER,
        spaceAfter=20,
    )

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=24,
        spaceBefore=12,
    )

    body_style = ParagraphStyle(
        "DocBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=16.5,  # 1.5 line spacing
        alignment=TA_LEFT,
        spaceAfter=6,
    )

    footer_style = ParagraphStyle(
        "DocFooter",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=7,
        textColor="#999999",
        alignment=TA_CENTER,
        spaceBefore=30,
    )

    # Build document content
    story: list = []

    # Header
    story.append(Paragraph("NYAY SETU — AI Generated Legal Document", header_style))
    story.append(Spacer(1, 6))

    # Title
    story.append(Paragraph(response.title, title_style))
    story.append(Spacer(1, 12))

    # Body — split content into paragraphs and render each
    content_lines: List[str] = response.content.split("\n")
    for line in content_lines:
        stripped: str = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
        else:
            # Escape XML special chars for ReportLab
            safe_line: str = (
                stripped.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(safe_line, body_style))

    # Sources
    if response.sources:
        story.append(Spacer(1, 20))
        sources_text: str = "Sources: " + ", ".join(response.sources)
        source_style = ParagraphStyle(
            "Sources",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            textColor="#888888",
        )
        story.append(Paragraph(sources_text, source_style))

    # Footer / Disclaimer
    story.append(Spacer(1, 30))
    story.append(
        Paragraph(
            "This document is AI-generated and should be reviewed by a qualified lawyer",
            footer_style,
        )
    )

    doc.build(story)
    buffer.seek(0)
    return buffer


def _create_docx(response: GenerateResponse, petitioner_name: str) -> io.BytesIO:
    """Convert generated document text to a simple DOCX in-memory file using python-docx."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed. Run: pip install python-docx",
        )

    doc = DocxDocument()

    # Header
    p = doc.add_paragraph()
    run = p.add_run("NYAY SETU — AI Generated Legal Document")
    run.bold = True
    run.font.size = Pt(8)

    # Title
    doc.add_paragraph(response.title).bold = True

    # Body — add each non-empty line as a paragraph
    for line in response.content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        else:
            para = doc.add_paragraph(stripped)
            para.style.font.size = Pt(11)

    # Sources
    if response.sources:
        doc.add_paragraph("\nSources:")
        doc.add_paragraph(", ".join(response.sources))

    # Disclaimer footer
    doc.add_paragraph("\nThis document is AI-generated and should be reviewed by a qualified lawyer")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ══════════════════════════════════════════════════════════════════════════════
# POST /generate — Generate legal document text
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/generate", response_model=GenerateResponse)
async def generate_document(request: GenerateRequest) -> GenerateResponse:
    """
    Generate a legal document (affidavit, RTI, complaint, or notice).
    Retrieves context from FAISS and calls an LLM to draft the document.
    """
    return _generate_document(request)


# ══════════════════════════════════════════════════════════════════════════════
# POST /generate/pdf — Generate legal document as PDF
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/generate/pdf")
async def generate_document_pdf(request: GenerateRequest):
    """
    Generate a legal document and return it as a downloadable PDF file.
    """
    response: GenerateResponse = _generate_document(request)
    pdf_buffer: io.BytesIO = _create_pdf(response, request.fields.petitioner_name)

    filename: str = f"{request.doc_type}_{request.fields.petitioner_name.replace(' ', '_')}.pdf"

    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )


@router.post("/generate/docx")
async def generate_document_docx(request: GenerateRequest):
    """Generate a legal document and return it as a downloadable DOCX file."""
    response: GenerateResponse = _generate_document(request)
    docx_buffer: io.BytesIO = _create_docx(response, request.fields.petitioner_name)

    filename: str = f"{request.doc_type}_{request.fields.petitioner_name.replace(' ', '_')}.docx"

    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
        },
    )
